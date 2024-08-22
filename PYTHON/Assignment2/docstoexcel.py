import os
from docx import Document
import openpyxl

def wordToExcel(directory, outputFileName):
    data = []
    products = ["Parka", "Boots", "Snowshoes", "Climbing Rope", "Oxygen Tank", "Ice Pick", "Crampons"]

    # List all .docx files in the specified directory
    docxFiles = [os.path.join(directory, file) for file in os.listdir(directory) if file.endswith(".docx")]

    # Iterate through each .docx file found
    for docxFile in docxFiles:
        # Load the Word document
        doc = Document(docxFile)
        invoice_id = doc.paragraphs[0].text.strip()  # Get invoice ID from the first paragraph
        total_products = 0
        subtotal = 0.0
        tax = 0.0
        total = 0.0
                             
        # Iterate through each paragraph in the document
        for para in doc.paragraphs:
            text = para.text
            # Split each line by newline character and process
            for line in text.split("\n"):
                name, value, *rest = (line + ":").split(":")
                if name == "SUBTOTAL":
                    subtotal = float(value)
                elif name == "TOTAL":
                    total = float(value)
                elif name == "TAX":
                    tax = float(value)
                elif name in products:
                    total_products += int(value)

        # Append invoice data to the list
        data.append([invoice_id, total_products, subtotal, tax, total])

    # Create a new Excel workbook and select the active sheet
    wb = openpyxl.Workbook()
    ws = wb.active

    # Add column headers
    headers = ["Invoice ID", "Total Products", "Subtotal", "Tax", "Total"]
    ws.append(headers)

    # Write the data to the Excel sheet
    for row in data:
        ws.append(row)

    # Save the workbook to the specified file
    wb.save(outputFileName)

# Directory where the .docx files are located
directory = r"C:/Users/Esteban/Desktop/University/Second Semester/Document Python/PYTHON"

# Output Excel file name
outputFileName = r"C:/Users/Esteban/Desktop/University/Second Semester/Document Python/PYTHON/invoices_summary.xlsx"

# Process the Word documents and create the Excel file
wordToExcel(directory, outputFileName)



"""
    for para in doc.paragraphs:
        text = para.text
        for line in text.split("\n"):
            if line.startswith("SUBTOTAL:"):
                subtotal = float(line.split("SUBTOTAL:")[1].split("\n")[0])
            elif line.startswith("TAX:"):
                tax = float(line.split("TAX:")[1].split("\n")[0])
            elif line.startswith("TOTAL:"):
                total = float(line.split("TOTAL:")[1].split("\n")[0])
            else:
                for product in products:
                    if product in line:
                        _, quantity = line.split(":")
                        total_products += int(quantity.strip())
"""

