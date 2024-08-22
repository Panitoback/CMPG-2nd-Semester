from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT

'''------------------------------------ DOCUMENT CREATION AND BASICS OPERATIONS -----------------------------------'''

doc = Document()

with open("report.txt", 'w') as file:
    file.write("Hola mi amigo, esto es solo una prueba \n testeando el archivo")

with open("report.txt", 'r') as file:
    content = file.readlines()

# Add a heading to the document
run_title = doc.add_heading("Assignment 2", 0)
run_title2 = doc.add_heading("Second line", 0)
run_title.runs[0].bold = True
run_title.runs[0].font.size = Pt(60)
run_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

'''------------------------------------ ADDING AND FORMATING TEXT -----------------------------------'''

# Add a paragraph with multiple runs
paragraph = doc.add_paragraph()
run1 = paragraph.add_run("This is a bold text. ")
run1.bold = True
run1.font.name = 'Arial'
run1.font.color.rgb = RGBColor(255, 0, 0)  # Red

run2 = paragraph.add_run("This is an italic text. ")
run2.italic = True
run2.font.name = 'Times New Roman'
run2.font.color.rgb = RGBColor(0, 255, 0)  # Green

run3 = paragraph.add_run("This is an underlined text. ")
run3.underline = True
run3.font.name = 'Courier New'
run3.font.color.rgb = RGBColor(0, 0, 255)  # Blue

# Create a bullet list
bullet_list = doc.add_paragraph()
bullet_list.style = 'List Bullet'

item1 = bullet_list.add_run("First item with a different font style.")
item1.font.name = 'Comic Sans MS'
item1.font.size = Pt(12)

bullet_list = doc.add_paragraph()
bullet_list.style = 'List Bullet'

item2 = bullet_list.add_run("Second item with another font style.")
item2.font.name = 'Verdana'
item2.font.size = Pt(14)

bullet_list = doc.add_paragraph()
bullet_list.style = 'List Bullet'

item3 = bullet_list.add_run("Third item with yet another font style.")
item3.font.name = 'Calibri'
item3.font.size = Pt(16)

'''------------------------------------ TABLES AND IMAGES -----------------------------------'''

# Insert a table with custom dimensions (3x3)
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'

# Fill the table with sample data and apply different styles
data = [
    ["Header 1", "Header 2", "Header 3"],
    ["Row 1, Cell 1", "Row 1, Cell 2", "Row 1, Cell 3"],
    ["Row 2, Cell 1", "Row 2, Cell 2", "Row 2, Cell 3"]
]

for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        cell.text = data[i][j]
        if i == 0:  # Styles for the header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            shading_elm = OxmlElement('w:shd') #Create a new element called 'shd'
            shading_elm.set(qn('w:fill'), "D3D3D3")  # Gray background color
            cell._element.get_or_add_tcPr().append(shading_elm)
        else:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

'''------------------------------------ SECTION AND HEADER FOOTERS -----------------------------------'''

# Function to add headers and footers to a section
def add_header_footer(section, header_text, footer_text):
    # Header
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run(header_text)
    header_run.bold = True
    header_run.font.size = Pt(14)
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.text = footer_text

    # Page number
    page_num_run = footer_paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

# Section 1
section1 = doc.sections[0]
add_header_footer(section1, "Section 1 Header", "Section 1 Footer")

# Content of Section 1
doc.add_paragraph("Content of Section 1.")

# Create a new section
section2 = doc.add_section()
add_header_footer(section2, "Section 2 Header", "Section 2 Footer")

# Content of Section 2
doc.add_paragraph("Content of Section 2.")

# Create another new section
section3 = doc.add_section()
add_header_footer(section3, "Section 3 Header", "Section 3 Footer")

# Content of Section 3
doc.add_paragraph("Content of Section 3.")

# Save the document
doc.save("output.docx")

'''------------------------------------ STYLES AND TEMPLATES -----------------------------------'''

# Define a custom style
style = doc.styles.add_style('CustomStyle', 1)  # 1 is for paragraph style
style.font.name = 'Arial'
style.font.size = Pt(12)
style.font.color.rgb = RGBColor(0, 0, 255)  # Blue
style.paragraph_format.space_after = Pt(12)

# Apply the custom style to several paragraphs
paragraph1 = doc.add_paragraph("This is the first paragraph with the custom style.", style='CustomStyle')
paragraph2 = doc.add_paragraph("This is the second paragraph with the custom style.", style='CustomStyle')
paragraph3 = doc.add_paragraph("This is the third paragraph with the custom style.", style='CustomStyle')

# Save the document
doc.save('styled_document.docx')

# Second part assignment
# Open the document based on an existing template
template_path = 'template.docx'
doc = Document(template_path)

# Modify the content of the document based on the template
doc.add_heading('Title of the Document Based on Template', level=1)

paragraph = doc.add_paragraph('This paragraph was added to the document based on the template.')
paragraph.style = doc.styles['Normal']  # Apply the document's normal style

# Create a new section and add content
section = doc.add_section()
section_header = section.header
section_header_paragraph = section_header.paragraphs[0]
section_header_paragraph.text = 'Header of the New Section'

# Add content to the new section
doc.add_paragraph('Content of the new section.')

# Save the modified document
doc.save('document_from_template.docx')

'''------------------------------------ ADVANCE FORMATING -----------------------------------'''

# Add a title and several headings
doc.add_heading('Document Title', level=0)
doc.add_heading('Section 1', level=1)
doc.add_paragraph('Content of section 1.')
doc.add_heading('Subsection 1.1', level=2)
doc.add_paragraph('Content of subsection 1.1.')
doc.add_heading('Section 2', level=1)
doc.add_paragraph('Content of section 2.')

# Add the table of contents
paragraph = doc.add_paragraph()
run = paragraph.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)
run._r.append(fldChar3)

# Add a blank page (optional)
doc.add_page_break()

# Create a table with a custom style
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'

# Define a custom table style
custom_table_style = doc.styles.add_style('CustomTableStyle', WD_STYLE_TYPE.TABLE)
table_style = doc.styles['CustomTableStyle']
table_style.font.name = 'Arial'
table_style.font.size = Pt(10)
table_style.font.color.rgb = RGBColor(255, 255, 255)

# Add background color to the header cells
table_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
for cell in table.rows[0].cells:
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '4F81BD')  # Dark blue
    cell._element.get_or_add_tcPr().append(shading_elm)
    cell.paragraphs[0].runs[0].font.bold = True

# Add data to the table
data = [
    ["Header 1", "Header 2", "Header 3"],
    ["Row 1, Cell 1", "Row 1, Cell 2", "Row 1, Cell 3"],
    ["Row 2, Cell 1", "Row 2, Cell 2", "Row 2, Cell 3"]
]

for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        cell.text = data[i][j]
        cell.paragraphs[0].style = 'CustomTableStyle'

# Save the document
doc.save('styled_document_with_toc.docx')
